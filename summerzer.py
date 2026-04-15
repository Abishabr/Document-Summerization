import os
from xml.dom.minidom import Document

import nltk
from nltk.corpus.reader import documents
from nltk.tokenize import  sent_tokenize
nltk.download('punkt', quiet=True)


documents = {}
sentences = {}

# ====================== FUNCTION TO READ FILE ======================

def readFile(filePath, fileName):
    try :
        if fileName.endswith('.txt'):
            with open(filePath, 'r') as f:
                return f.read()
        elif fileName.endswith('.pdf'):
            from PyPDF2 import PdfReader
            reader = PdfReader(filePath)
            text = ""
            for page in reader.pages:
                text += page.extractText() or ""
            return text
        elif fileName.endswith('.docx'):
            from docx import Documet
            doc = Document(filePath)
            text = "/n".join([p.text for p in doc.paragraphs])
            return text
        return ""
    except Exception as e:
        print(e)

# ====================== MAIN FUNCTION - LOAD DOCUMENTS ======================
def loadDocument(folder_name = "Documents"):
    print("Starting to load documents...\n")
    if not os.path.exists(folder_name):
        print("Folder doesn't exist.")
        return
    count = 0

    for fileName in os.listdir(folder_name):
        filePath = os.path.join(folder_name, fileName)


        if fileName.endswith(('.txt', '.pdf', '.docx')):
            text = readFile(filePath, fileName)
            if text and len(text.strip()) > 100:
                count += 1
                docId = f"{count}"
                documents[docId] = text
                sentences[docId] = sent_tokenize(text)

                print(f"✅ Loaded: {fileName}   ({len(sentences[docId])} sentences)")
            else
                print(f"⚠️ Skipped: {fileName} (too short or empty)")
    print(f"\n🎉 Successfully loaded {count} document(s)!")

# ====================== SHOW STATUS ======================
def showStatus():
    print("/n === Current Status ====")
    print(f"Total documents loaded: {len(documents)}")

    for docId in documents:
        print(f"   • {docId} → {len(sentences[docId])} sentences")

if __name__ == "__main__":
    loadDocument("documents")   # Change folder name if needed
    showStatus()

